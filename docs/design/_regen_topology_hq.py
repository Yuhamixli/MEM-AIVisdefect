# -*- coding: utf-8 -*-
"""Regenerate Canvas-quality software topology PNG and swap into docx (Track Changes)."""
from __future__ import annotations

import copy
import re
import shutil
from datetime import datetime, timezone
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu
from lxml import etree
from playwright.sync_api import sync_playwright

DOCX = Path(
    r"a:\Projects\MEM-AIVisdefect\docs\design"
    r"\03组_玻纤拉挤电芯压条_AI视觉检测_详细方案设计_编号目录整理版_算法软件修订.docx"
)
FIG_DIR = Path(r"a:\Projects\MEM-AIVisdefect\docs\design\_figs")
FIG_PATH = FIG_DIR / "software-topology.png"
FIG_OLD_BAK = FIG_DIR / "software-topology.before-hq.png"
HTML_PATH = FIG_DIR / "software-topology.html"
AUTHOR = "算法软件组"
DATE = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

# Mirror canvas nodes / edges
NODES = [
    {"id": "line", "label": "五段线 + PLC", "sub": "机械/触发 · 演进", "kind": "evo"},
    {"id": "optics", "label": "受控成像", "sub": "6MP 卷帘 + DOE 光源", "kind": "evo"},
    {"id": "raw", "label": "raw-data", "sub": "版本化原图 · 文件事实源", "kind": "data"},
    {"id": "l1", "label": "在线 L1", "sub": "PatchCore → 可选 YOLO", "kind": "evo"},
    {"id": "near", "label": "近线精修", "sub": "DINO + SAM2 工具层", "kind": "evo"},
    {"id": "offline", "label": "离线检测模块", "sub": "yolo11s-seg · P0 硬交付", "kind": "p0"},
    {"id": "flywheel", "label": "数据飞轮 L4", "sub": "VLM → 人工确认 → 反哺", "kind": "evo"},
    {"id": "out", "label": "result.json", "sub": "+ overlay.jpg", "kind": "data"},
    {"id": "ui", "label": "detector-ui", "sub": "结果 / 复核", "kind": "app"},
    {"id": "bi", "label": "BI 看板", "sub": "Cloudflare Pages", "kind": "app"},
    {"id": "bridge", "label": "飞书桥", "sub": "协作侧车 · 非 AOI", "kind": "collab"},
]
EDGES = [
    ("line", "optics"),
    ("optics", "raw"),
    ("raw", "l1"),
    ("raw", "offline"),
    ("l1", "near"),
    ("near", "flywheel"),
    ("offline", "out"),
    ("flywheel", "offline"),  # back-edge
    ("out", "ui"),
    ("raw", "bi"),
    ("bridge", "bi"),
]

KIND_STYLE = {
    "p0": {
        "fill": "#DCEFE4",
        "stroke": "#178A4B",
        "strokeWidth": 2.75,
        "dash": None,
        "label": "P0 交付",
    },
    "evo": {
        "fill": "#FBFBFC",
        "stroke": "#8B919A",
        "strokeWidth": 1.35,
        "dash": "5 3.5",
        "label": "演进 / 虚线",
    },
    "data": {
        "fill": "#E8F1F8",
        "stroke": "#3A6F9A",
        "strokeWidth": 1.5,
        "dash": None,
        "label": "数据层",
    },
    "app": {
        "fill": "#FFF3DE",
        "stroke": "#B07A00",
        "strokeWidth": 1.5,
        "dash": None,
        "label": "应用层",
    },
    "collab": {
        "fill": "#F3F3F5",
        "stroke": "#8B919A",
        "strokeWidth": 1.35,
        "dash": "5 3.5",
        "label": "协作侧车",
    },
}


def compute_dag_layout(
    nodes: list[dict],
    edges: list[tuple[str, str]],
    *,
    node_w: float = 210,
    node_h: float = 66,
    rank_gap: float = 78,
    node_gap: float = 40,
    padding: float = 32,
) -> dict:
    """Sugiyama-style vertical ranks (Canvas computeDAGLayout-inspired)."""
    ids = [n["id"] for n in nodes]
    idset = set(ids)
    fwd = [(a, b) for a, b in edges if a in idset and b in idset]
    # detect back-edges via DFS on undirected cycle of feedback: edges that go to already visited in topo attempt
    # Simple: Kahn topo; leftover edges that violate are back-edges
    succ: dict[str, list[str]] = {i: [] for i in ids}
    pred_count = {i: 0 for i in ids}
    for a, b in fwd:
        succ[a].append(b)
        pred_count[b] += 1

    # Find back-edges: temporarily ignore edges that create cycles
    # Use iterative removal of cycle edges: if node has self-loop path
    back = set()
    # Prefer marking flywheel->offline as back if present
    if ("flywheel", "offline") in fwd:
        back.add(("flywheel", "offline"))
        pred_count["offline"] -= 1
        succ["flywheel"] = [t for t in succ["flywheel"] if t != "offline"]

    # ranks via longest-path from sources
    rank: dict[str, int] = {}
    from collections import deque

    q = deque([i for i in ids if pred_count[i] == 0])
    for i in ids:
        if pred_count[i] == 0:
            rank[i] = 0
    seen = set(q)
    while q:
        u = q.popleft()
        for v in succ[u]:
            rank[v] = max(rank.get(v, 0), rank[u] + 1)
            pred_count[v] -= 1
            if pred_count[v] == 0 and v not in seen:
                seen.add(v)
                q.append(v)
    # any remaining (shouldn't) get max+1
    max_r = max(rank.values()) if rank else 0
    for i in ids:
        rank.setdefault(i, max_r)

    # group by rank
    by_rank: dict[int, list[str]] = {}
    for i in ids:
        by_rank.setdefault(rank[i], []).append(i)

    # order within rank: prefer left-to-right by barycenter of parents
    parents: dict[str, list[str]] = {i: [] for i in ids}
    for a, b in fwd:
        if (a, b) in back:
            continue
        parents[b].append(a)

    # First pass: order ranks; provisional x-centers for barycenter (use max width)
    max_nodes = max(len(g) for g in by_rank.values()) if by_rank else 1
    content_w = max_nodes * node_w + (max_nodes - 1) * node_gap
    width = content_w + padding * 2

    ordered_ranks: list[list[str]] = []
    pos_x_center: dict[str, float] = {}
    for r in sorted(by_rank):
        group = by_rank[r]
        preferred_orders = {
            0: ["line", "bridge"],
        }
        # Stable visual order for key ranks (Canvas-like reading order)
        force_order = {
            "l1": 0,
            "offline": 1,
            "bi": 2,
            "near": 0,
            "out": 1,
            "flywheel": 0,
            "ui": 1,
        }
        if r in preferred_orders:
            preferred = preferred_orders[r]
            group = sorted(
                group,
                key=lambda x: preferred.index(x) if x in preferred else 50 + ids.index(x),
            )
        elif any(n in force_order for n in group):
            group = sorted(
                group,
                key=lambda x: force_order.get(x, 50 + ids.index(x)),
            )
        else:

            def key(nid: str) -> float:
                ps = parents[nid]
                if not ps:
                    return float(ids.index(nid))
                return sum(pos_x_center[p] for p in ps) / len(ps)

            group = sorted(group, key=key)
        ordered_ranks.append(group)
        n = len(group)
        total_w = n * node_w + (n - 1) * node_gap
        start_x = (width - total_w) / 2
        for j, nid in enumerate(group):
            pos_x_center[nid] = start_x + j * (node_w + node_gap) + node_w / 2

    height = len(ordered_ranks) * node_h + (len(ordered_ranks) - 1) * rank_gap + padding * 2

    layout_nodes = []
    ranks_meta = []
    for ri, group in enumerate(ordered_ranks):
        n = len(group)
        total_w = n * node_w + (n - 1) * node_gap
        start_x = (width - total_w) / 2
        y = padding + ri * (node_h + rank_gap)
        ranks_meta.append(
            {
                "rank": ri,
                "x": padding * 0.45,
                "y": y - 14,
                "width": width - padding * 0.9,
                "height": node_h + 28,
            }
        )
        for j, nid in enumerate(group):
            x = start_x + j * (node_w + node_gap)
            pos_x_center[nid] = x + node_w / 2
            layout_nodes.append(
                {
                    "id": nid,
                    "x": x,
                    "y": y,
                    "cx": x + node_w / 2,
                    "cy": y + node_h / 2,
                    "bottom": y + node_h,
                    "top": y,
                }
            )

    by_id = {n["id"]: n for n in layout_nodes}
    layout_edges = []
    for a, b in fwd:
        sa, sb = by_id[a], by_id[b]
        is_back = (a, b) in back
        layout_edges.append(
            {
                "from": a,
                "to": b,
                "sourceX": sa["cx"],
                "sourceY": sa["bottom"] if not is_back else sa["top"],
                "targetX": sb["cx"],
                "targetY": sb["top"] if not is_back else sb["bottom"],
                "isBackEdge": is_back,
            }
        )

    return {
        "width": width,
        "height": height,
        "nodeW": node_w,
        "nodeH": node_h,
        "nodes": layout_nodes,
        "edges": layout_edges,
        "ranks": ranks_meta,
    }


def build_html(layout: dict) -> str:
    meta = {n["id"]: n for n in NODES}
    node_w = layout["nodeW"]
    node_h = layout["nodeH"]
    w = layout["width"]
    h = layout["height"]

    band_rects = []
    for r in layout["ranks"]:
        band_rects.append(
            f'<rect x="{r["x"]:.1f}" y="{r["y"]:.1f}" width="{r["width"]:.1f}" '
            f'height="{r["height"]:.1f}" rx="8" fill="#F0F2F5" opacity="0.55"/>'
        )

    edge_paths = []
    for e in layout["edges"]:
        sx, sy, tx, ty = e["sourceX"], e["sourceY"], e["targetX"], e["targetY"]
        if e["isBackEdge"]:
            # Feedback loop on the left side (flywheel → offline)
            mid_x = min(sx, tx) - 56
            d = (
                f"M {sx:.1f} {sy:.1f} C {mid_x:.1f} {sy:.1f}, "
                f"{mid_x:.1f} {ty:.1f}, {tx:.1f} {ty:.1f}"
            )
            stroke = "#178A4B"
            dash = ' stroke-dasharray="6 4"'
            opacity = "0.92"
            marker = ' marker-end="url(#arrowBack)"'
        else:
            mid_y = (sy + ty) / 2
            d = (
                f"M {sx:.1f} {sy:.1f} C {sx:.1f} {mid_y:.1f}, "
                f"{tx:.1f} {mid_y:.1f}, {tx:.1f} {ty:.1f}"
            )
            stroke = "#6A7180"
            dash = ""
            opacity = "0.78"
            marker = ' marker-end="url(#arrowFwd)"'
        edge_paths.append(
            f'<path d="{d}" fill="none" stroke="{stroke}" stroke-width="1.8"'
            f'{dash} opacity="{opacity}"{marker}/>'
        )

    node_gs = []
    for n in layout["nodes"]:
        m = meta[n["id"]]
        st = KIND_STYLE[m["kind"]]
        dash_attr = f' stroke-dasharray="{st["dash"]}"' if st["dash"] else ""
        node_gs.append(
            f'''<g transform="translate({n["x"]:.1f},{n["y"]:.1f})">
  <rect width="{node_w}" height="{node_h}" rx="8" fill="{st["fill"]}"
        stroke="{st["stroke"]}" stroke-width="{st["strokeWidth"]}"{dash_attr}/>
  <text x="{node_w/2}" y="27" text-anchor="middle" fill="#151820"
        font-size="15.5" font-weight="650" font-family="Microsoft YaHei, DengXian, SimHei, sans-serif">{m["label"]}</text>
  <text x="{node_w/2}" y="48" text-anchor="middle" fill="#5C6370"
        font-size="12.5" font-family="Microsoft YaHei, DengXian, SimHei, sans-serif">{m["sub"]}</text>
</g>'''
        )

    legend_items = []
    x0 = 0
    for kind in ("p0", "evo", "data", "app", "collab"):
        st = KIND_STYLE[kind]
        dash_attr = f' stroke-dasharray="{st["dash"]}"' if st["dash"] else ""
        legend_items.append(
            f'''<g transform="translate({x0},0)">
  <rect width="18" height="14" rx="3" y="1" fill="{st["fill"]}" stroke="{st["stroke"]}"
        stroke-width="1.5"{dash_attr}/>
  <text x="24" y="13" font-size="12" fill="#3A3F48"
        font-family="Microsoft YaHei, DengXian, SimHei, sans-serif">{st["label"]}</text>
</g>'''
        )
        x0 += 118

    # outer chrome: title + subtitle + legend + svg
    outer_w = max(w + 80, 1180)
    svg_x = (outer_w - w) / 2
    legend_x = (outer_w - x0) / 2
    total_h = h + 168

    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8"/>
<style>
  html, body {{
    margin: 0; padding: 0;
    background: #FFFFFF;
    -webkit-font-smoothing: antialiased;
  }}
  #root {{
    width: {outer_w}px;
    padding: 36px 40px 40px;
    box-sizing: border-box;
    font-family: "Microsoft YaHei", "DengXian", "SimHei", sans-serif;
    color: #1A1D23;
  }}
  h1 {{
    margin: 0 0 8px;
    font-size: 28px;
    font-weight: 700;
    letter-spacing: 0.02em;
  }}
  .sub {{
    margin: 0 0 18px;
    font-size: 13.5px;
    color: #5C6370;
    line-height: 1.45;
  }}
  .legend {{
    margin-bottom: 16px;
  }}
  .frame {{
    border: 1px solid #E2E5EA;
    border-radius: 12px;
    background: #FFFFFF;
    padding: 12px 8px 8px;
  }}
</style>
</head>
<body>
<div id="root">
  <h1>全项目软件拓扑图</h1>
  <p class="sub">MEM-AIVisdefect · 拉挤表面 AOI · 文件事实源 · 在线/近线/离线拆分（ADR-001）· 实心强调 = P0 硬交付</p>
  <svg class="legend" width="{outer_w - 80}" height="22" viewBox="0 0 {outer_w - 80} 22">
    <g transform="translate({legend_x - 40:.1f},0)">
      {"".join(legend_items)}
    </g>
  </svg>
  <div class="frame">
    <svg width="{w}" height="{h}" viewBox="0 0 {w} {h}" role="img" aria-label="Software topology DAG"
         style="display:block;margin:0 auto;">
      <defs>
        <marker id="arrowFwd" viewBox="0 0 10 10" refX="8" refY="5" markerWidth="7" markerHeight="7" orient="auto-start-reverse">
          <path d="M 0 1.2 L 8.5 5 L 0 8.8" fill="none" stroke="#6A7180" stroke-width="1.4" stroke-linecap="round" stroke-linejoin="round"/>
        </marker>
        <marker id="arrowBack" viewBox="0 0 10 10" refX="8" refY="5" markerWidth="7" markerHeight="7" orient="auto-start-reverse">
          <path d="M 0 1.2 L 8.5 5 L 0 8.8" fill="none" stroke="#178A4B" stroke-width="1.4" stroke-linecap="round" stroke-linejoin="round"/>
        </marker>
      </defs>
      {"".join(band_rects)}
      {"".join(edge_paths)}
      {"".join(node_gs)}
    </svg>
  </div>
</div>
</body>
</html>
"""


def render_png() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    if FIG_PATH.exists() and not FIG_OLD_BAK.exists():
        shutil.copy2(FIG_PATH, FIG_OLD_BAK)

    layout = compute_dag_layout(NODES, EDGES)
    html = build_html(layout)
    HTML_PATH.write_text(html, encoding="utf-8")

    with sync_playwright() as p:
        browser = p.chromium.launch()
        page = browser.new_page(
            viewport={"width": 1400, "height": 1200},
            device_scale_factor=2.5,
        )
        page.goto(HTML_PATH.resolve().as_uri(), wait_until="networkidle")
        root = page.locator("#root")
        root.screenshot(path=str(FIG_PATH), type="png")
        browser.close()

    # Ensure high-res width: if under 2400, upscale with LANCZOS
    from PIL import Image

    im = Image.open(FIG_PATH)
    if im.width < 2400:
        scale = 2600 / im.width
        im = im.resize((int(im.width * scale), int(im.height * scale)), Image.Resampling.LANCZOS)
        im.save(FIG_PATH, "PNG", optimize=True)
    print(f"PNG: {FIG_PATH} size={im.size} bytes={FIG_PATH.stat().st_size}")


def enable_track_revisions(doc: Document) -> None:
    settings = doc.settings.element
    if settings.find(qn("w:trackRevisions")) is None:
        settings.insert(0, OxmlElement("w:trackRevisions"))


def next_revision_id(doc: Document) -> int:
    body = doc.element.body
    ids = []
    for tag in ("w:ins", "w:del"):
        for el in body.findall(".//" + qn(tag)):
            rid = el.get(qn("w:id"))
            if rid is not None and str(rid).isdigit():
                ids.append(int(rid))
    return (max(ids) + 1) if ids else 1


def paragraph_plain_text(p_el: etree._Element) -> str:
    parts = []
    for node in p_el.iter():
        if node.tag == qn("w:delText"):
            continue
        if node.tag == qn("w:t") and node.text:
            parent = node.getparent()
            skip = False
            while parent is not None:
                if parent.tag == qn("w:del"):
                    skip = True
                    break
                parent = parent.getparent()
            if not skip:
                parts.append(node.text)
    return "".join(parts)


def find_topology_picture_paragraph(doc: Document):
    """Find the paragraph containing the topology figure (before 图1 caption or after it)."""
    paras = list(doc.paragraphs)
    cap_idx = None
    for i, p in enumerate(paras):
        plain = paragraph_plain_text(p._element)
        if "全项目软件拓扑" in plain and ("图" in plain or "图 1" in plain or "图1" in plain):
            cap_idx = i
            break
        if plain.strip() in ("图 1  全项目软件拓扑图", "图1  全项目软件拓扑图"):
            cap_idx = i
            break

    # Prefer picture immediately after caption, else before
    candidates = []
    if cap_idx is not None:
        for j in (cap_idx + 1, cap_idx - 1, cap_idx + 2):
            if 0 <= j < len(paras):
                candidates.append(paras[j])
    for p in candidates + paras:
        xml = etree.tostring(p._element, encoding="unicode")
        if "a:blip" in xml or "w:drawing" in xml or "v:imagedata" in xml:
            # Prefer the one near caption
            if cap_idx is not None:
                # check proximity
                idx = paras.index(p) if p in paras else -1
                if idx >= 0 and abs(idx - cap_idx) <= 3:
                    return p, cap_idx
            else:
                return p, None
    # fallback: any blip in section 3
    for i, p in enumerate(paras):
        if "a:blip" in etree.tostring(p._element, encoding="unicode"):
            return p, cap_idx
    return None, cap_idx


def replace_picture_with_tracked_ins(
    doc: Document, pic_para, fig_path: Path, rev_id_start: int
) -> int:
    """
    Tracked image replacement strategy:
    1. Wrap existing drawing run(s) in w:del (clone into del).
    2. Insert new picture run wrapped in w:ins after the del.
    Returns next revision id.
    """
    rid = rev_id_start
    p_el = pic_para._element

    # Collect drawing-bearing runs
    drawing_runs = []
    for child in list(p_el):
        tag = child.tag.split("}")[-1]
        if tag == "pPr":
            continue
        xml = etree.tostring(child, encoding="unicode")
        if "w:drawing" in xml or "a:blip" in xml or "v:imagedata" in xml or "w:pict" in xml:
            drawing_runs.append(child)

    if not drawing_runs:
        raise RuntimeError("No drawing runs found in picture paragraph")

    # Build w:del containing clones of old drawing runs
    dele = OxmlElement("w:del")
    dele.set(qn("w:id"), str(rid))
    rid += 1
    dele.set(qn("w:author"), AUTHOR)
    dele.set(qn("w:date"), DATE)
    for run_el in drawing_runs:
        clone = copy.deepcopy(run_el)
        # Word expects deleted runs may need w:delText for text; drawings stay as-is inside w:del
        dele.append(clone)
        p_el.remove(run_el)

    # Insert del at end of paragraph (keep pPr)
    p_el.append(dele)

    # Add new picture via python-docx into a temporary paragraph, then move into w:ins
    tmp = doc.add_paragraph()
    tmp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tmp.add_run()
    run.add_picture(str(fig_path), width=Cm(16.2))
    new_run_el = tmp.runs[0]._element

    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), str(rid))
    rid += 1
    ins.set(qn("w:author"), AUTHOR)
    ins.set(qn("w:date"), DATE)
    ins.append(copy.deepcopy(new_run_el))
    p_el.append(ins)

    # Remove temp paragraph
    tmp._element.getparent().remove(tmp._element)

    # Ensure paragraph centered
    pPr = p_el.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_el.insert(0, pPr)
    jc = pPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        pPr.append(jc)
    jc.set(qn("w:val"), "center")

    return rid


def add_tracked_note_after(doc: Document, ref_el, text: str, rev_id: int) -> int:
    p = doc.add_paragraph()
    p_el = p._element
    # clear
    for child in list(p_el):
        if child.tag != qn("w:pPr"):
            p_el.remove(child)
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), str(rev_id))
    rev_id += 1
    ins.set(qn("w:author"), AUTHOR)
    ins.set(qn("w:date"), DATE)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "666666")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")
    rpr.append(color)
    rpr.append(sz)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    ins.append(run)
    p_el.append(ins)
    ref_el.addnext(p_el)
    return rev_id


def update_docx() -> None:
    bak = DOCX.with_suffix(".docx.bak-before-hq-topology")
    if not bak.exists():
        bak.write_bytes(DOCX.read_bytes())
        print("Backup:", bak)
    else:
        # always refresh a working bak timestamped? keep first; also write overwrite bak
        bak.write_bytes(DOCX.read_bytes())
        print("Backup refreshed:", bak)

    doc = Document(str(DOCX))
    enable_track_revisions(doc)
    rid = next_revision_id(doc)

    pic_para, cap_idx = find_topology_picture_paragraph(doc)
    if pic_para is None:
        raise RuntimeError("Could not locate topology picture paragraph in docx")

    print(
        "Found picture para near caption idx=",
        cap_idx,
        "plain=",
        paragraph_plain_text(pic_para._element)[:40],
    )
    rid = replace_picture_with_tracked_ins(doc, pic_para, FIG_PATH, rid)

    # Optional short tracked caption note if caption exists
    if cap_idx is not None:
        cap = doc.paragraphs[cap_idx]
        # Append a tiny revision note after the picture paragraph
        note = (
            "【修订·算法软件组】拓扑图已替换为与 Canvas 一致的 DAG 版（分层色带 / P0 实线强调 / 演进虚线框）。"
        )
        # Avoid duplicating if already present
        nxt = pic_para._element.getnext()
        already = False
        if nxt is not None:
            txt = "".join(nxt.itertext())
            if "拓扑图已替换为与 Canvas 一致" in txt:
                already = True
        if not already:
            rid = add_tracked_note_after(doc, pic_para._element, note, rid)

    doc.save(str(DOCX))
    print("Saved", DOCX)

    # Verify
    doc2 = Document(str(DOCX))
    tr = doc2.settings.element.find(qn("w:trackRevisions"))
    print("trackRevisions:", tr is not None)
    body = doc2.element.body
    print("w:ins:", len(body.findall(".//" + qn("w:ins"))))
    print("w:del:", len(body.findall(".//" + qn("w:del"))))
    # Confirm new media references PNG size
    with ZipFile(DOCX) as z:
        medias = sorted(n for n in z.namelist() if n.startswith("word/media/"))
        for m in medias:
            info = z.getinfo(m)
            print(f"  {m}: {info.file_size} bytes")


def patch_docx_media_only() -> None:
    """Replace the already-inserted HQ image part without nesting more revisions."""
    import tempfile
    import zipfile

    target_name = None
    with ZipFile(DOCX, "r") as z:
        # Prefer the largest / newest topology media (image2 after tracked insert)
        medias = sorted(n for n in z.namelist() if n.startswith("word/media/") and n.endswith(".png"))
        if not medias:
            raise RuntimeError("no media png in docx")
        # Choose media matching current FIG_PATH size closely, else last media
        fig_size = FIG_PATH.stat().st_size
        best = None
        best_diff = 10**18
        for name in medias:
            diff = abs(z.getinfo(name).file_size - fig_size)
            # also allow previous HQ size
            if diff < best_diff:
                best_diff = diff
                best = name
        # After first tracked insert, image2 is the new one; image1 is old matplotlib
        target_name = "word/media/image2.png" if "word/media/image2.png" in medias else best
        print(f"Patching media part: {target_name}")

    buf = FIG_PATH.read_bytes()
    tmp = DOCX.with_suffix(".docx.tmp-media")
    with ZipFile(DOCX, "r") as zin, ZipFile(tmp, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == target_name:
                data = buf
            zout.writestr(item, data)
    tmp.replace(DOCX)
    print("Patched", DOCX, "media bytes", len(buf))


def main() -> None:
    import sys

    mode = sys.argv[1] if len(sys.argv) > 1 else "all"
    print("=== Render HQ topology PNG ===")
    render_png()
    if mode == "png":
        print("PNG-only done.")
        return
    if mode == "patch-media":
        print("=== Patch docx media (keep revisions) ===")
        patch_docx_media_only()
    else:
        print("=== Update docx under Track Changes ===")
        update_docx()
    print("Done.")


if __name__ == "__main__":
    main()
