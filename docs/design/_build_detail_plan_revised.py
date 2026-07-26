# -*- coding: utf-8 -*-
"""
1) Build a revised detail-design docx (python-docx, no track marks)
2) Word.CompareDocuments(original, revised) -> tracked-changes docx
"""
from __future__ import annotations

import shutil
import sys
import time
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

ORIG = Path(
    r"A:\Backup\xwechat_files\yuhan4988_a056\msg\file\2026-07"
    r"\03组_玻纤拉挤电芯压条_AI视觉检测_详细方案设计_编号目录整理版(1).docx"
)
OUT_DIR = Path(r"A:\Projects\MEM-AIVisdefect\docs\design")
TMP_DIR = Path(r"A:\Projects\MEM-AIVisdefect\.tmp")
REVISED_CLEAN = TMP_DIR / "detail_plan_revised_clean.docx"
OUT_TRACKED = OUT_DIR / "03组_玻纤拉挤电芯压条_AI视觉检测_详细方案设计_编号目录整理版_算法软件修订.docx"


def iter_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p


def para_text(p: Paragraph) -> str:
    return (p.text or "").strip()


def insert_paragraph_after(paragraph: Paragraph, text: str, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    if text:
        new_para.add_run(text)
    return new_para


def insert_paragraph_before(paragraph: Paragraph, text: str, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    if text:
        new_para.add_run(text)
    return new_para


def find_para(doc: Document, exact: str | None = None, startswith: str | None = None, in_toc: bool = False):
    for p in doc.paragraphs:
        t = para_text(p)
        style = p.style.name if p.style is not None else ""
        if in_toc and "toc" not in style.lower():
            continue
        if not in_toc and "toc" in style.lower():
            continue
        if exact is not None and t == exact:
            return p
        if startswith is not None and t.startswith(startswith):
            return p
    return None


def replace_para_text(p: Paragraph, new_text: str):
    # keep style; replace runs
    if p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(new_text)


def build_revised_clean() -> None:
    seed = TMP_DIR / f"_seed_{int(time.time())}.docx"
    shutil.copy2(ORIG, seed)
    doc = Document(str(seed))

    # --- renames in body + toc ---
    rename_map = [
        ("8 离线检测模块与接口设计", "8 离线检测模块与软件接口"),
        ("9 测试情况与效率对比", "10 测试情况与效率对比"),
        ("9.1 测试情况", "10.1 测试情况"),
        ("9.2 效率对比", "10.2 效率对比"),
    ]
    for p in doc.paragraphs:
        t = para_text(p)
        for old, new in rename_map:
            if t == old or t.startswith(old + "\t"):
                # preserve trailing tab+page if present
                suffix = ""
                if "\t" in (p.text or ""):
                    suffix = (p.text or "")[len(old) :]
                    # if text had no exact old at start after strip issues:
                raw = p.text or ""
                if raw.startswith(old):
                    replace_para_text(p, new + raw[len(old) :])
                else:
                    replace_para_text(p, new)
                break

    # --- Chapter 3 ---
    h3 = find_para(doc, exact="3 总体方案设计")
    if not h3:
        raise RuntimeError("ch3 heading missing")
    p = insert_paragraph_after(h3, "3.1 总体架构原则", "Heading 2")
    p = insert_paragraph_after(
        p,
        "验收主线尽量短：任务书P0指标由离线YOLO检测模块直接承担。"
        "离线主检、近线精修与数据飞轮分开部署，禁止叠成一条串行唯一判决链。"
        "受控成像先于算法：光源和拍摄规范经试验后冻结，成像不定则指标不可复现。",
        "Normal",
    )
    p = insert_paragraph_after(p, "3.2 软件能力边界总览", "Heading 2")
    p = insert_paragraph_after(
        p,
        "本期软件能力按数据流划分为四段："
        "（1）成像与原始数据入库（第5、6章，非本章实现重点）；"
        "（2）离线检测模块完成缺陷定位/分类并落盘结构化结果（第7、8章）；"
        "（3）同步脚本将算法落盘结果转换为前端可消费的jobs数据；"
        "（4）detector-ui承担检测结果展示与复核，管理BI承担进度/风险/M币/知识库/意见箱等管理信息展示（第9章）。"
        "在线产线节拍闭环与PLC联动不作为本期硬交付。",
        "Normal",
    )
    insert_paragraph_after(
        p,
        "逻辑关系可概括为：原图入库→与训练一致的预处理→YOLO主检（可选并联EfficientAD异常摘要）"
        "→result.json与overlay→sync转换→detector-ui/BI消费。"
        "近线工具（开放词汇定位、难例分割、多模态预标注）仅服务难例精修与样本回灌，"
        "不进入离线模块打包依赖，不承担实时放行承诺。",
        "Normal",
    )

    # --- Chapter 7 enrich ---
    p71 = find_para(doc, startswith="YOLO类轻量监督检测器")
    if not p71:
        raise RuntimeError("7.1 body missing")
    p = insert_paragraph_after(
        p71,
        "补充口径（对齐《拉挤表面缺陷离线检测与近线精修实施方案》）："
        "YOLO族有监督检测器为本期验收主检，承担已知缺陷定位与类别判定；"
        "EfficientAD/PatchCore等异常筛查仅作可选并联，用于补漏与难例发现，"
        "异常分支不得删除YOLO已检出结果，也不进入P0查全率分母的串行折算。"
        "基线模型采用YOLO11-seg（可退化输出检测框，最终定版以选型关单为准）。",
        "Normal",
    )
    p = insert_paragraph_after(p, "7.3 近线精修与工具边界", "Heading 2")
    insert_paragraph_after(
        p,
        "近线精修用于复核、难例处理和数据生产，不挡P0验收。"
        "开放词汇定位（如Grounding DINO）用于少样本探查与提示框生成；"
        "SAM2按需提供不规则缺陷掩膜；VLM仅做预标注建议，入库前须人工确认（建议抽检不低于10%）。"
        "禁则：不得以近线模型顶替离线模块作为第三方交付主程序；"
        "不得宣称大模型异步分析可实时防止漏检；"
        "未经人工确认的自动预标注不得直接计入金标准。",
        "Normal",
    )

    # --- Chapter 8 enrich ---
    h81 = find_para(doc, exact="8.1 功能模块")
    if not h81:
        raise RuntimeError("8.1 missing")
    insert_paragraph_after(
        h81,
        "离线检测模块是任务书“缺陷定义卡+检测模型+离线检测模块”的软件载体，"
        "要求可独立运行、可无网演示；第三方按使用说明完成“放图→检测→查看结构化结果/叠加图”。"
        "主路径输出缺陷类别、像素坐标框、可选掩膜、置信度、模型版本、推理耗时与初判结论。"
        "判定初版：无缺陷为OK；检出已知类缺陷默认REVIEW；NG规则待允收标准共签后启用。"
        "输出为目检初筛辅助，不替代最终放行。",
        "Normal",
    )

    h82 = find_para(doc, exact="8.2 结构化结果字段")
    if h82:
        insert_paragraph_after(
            h82,
            "逐图输出result.json，建议字段包括：schema_version、status、piece_id、batch_id、surface、"
            "image_path、image_size、model_version、pipeline、inference_time_ms、timestamp、"
            "defects[]、summary，以及可选anomaly（EfficientAD启用时）。"
            "defects[]含defect_id、class_slug、class_name、bbox{x,y,w,h}、confidence、可选mask；"
            "summary含total_defects、by_class、overall_decision与review_status。"
            "机器初判与人工复核分列，复核不得回写overall_decision，以便误报漏报分析与50件盲测。",
            "Normal",
        )

    h83 = find_para(doc, exact="8.3 接口定义")
    if h83:
        insert_paragraph_after(
            h83,
            "算法侧以CLI为主：detect --input --output --model [--conf] [--device] [--save-overlay] "
            "[--config] [--timeout-ms] [--anomaly-model]。"
            "配置文件与命令行参数对应，命令行优先。"
            "退出码：0全部成功，1部分失败，2致命错误。"
            "输出目录建议按{date}/{piece_id}/落盘image、result.json、overlay（及可选heatmap）。"
            "字段级细则以仓库《离线检测模块接口设计方案》为准。",
            "Normal",
        )

    h84 = find_para(doc, exact="8.4 部署与运行环境")
    if h84:
        p = insert_paragraph_after(
            h84,
            "运行环境基线：Python 3.10、Ultralytics、PyTorch；默认CPU可跑为硬约束，GPU为加速可选项。"
            "交付包含入口程序、默认配置、权重、样例图与接口说明。"
            "离线模块不依赖在线API与云端推理。",
            "Normal",
        )
        p = insert_paragraph_after(p, "8.5 与前端的数据交接", "Heading 2")
        insert_paragraph_after(
            p,
            "算法侧只写约定目录下的result.json、图像副本与叠加图；"
            "由同步脚本（sync-results/sync-data）转换为detector-ui可读的"
            "/data/detect/jobs/{piece_id}.json与jobs-index.json。"
            "算法侧与前端侧均不得绕过转换私改对方格式。"
            "字段映射要点：class_slug→slug，bbox对象→[x,y,w,h]，surface→faces[].face，"
            "summary.review_status→件级review_status。契约见第9章与《22-前后端API契约》。",
            "Normal",
        )

    # --- New chapter 9 before chapter 10 ---
    h10 = find_para(doc, exact="10 测试情况与效率对比")
    if not h10:
        raise RuntimeError("ch10 heading missing after rename")

    blocks = [
        ("Heading 1", "9 检测软件与BI详细设计"),
        ("Heading 2", "9.1 检测结果展示（detector-ui）"),
        (
            "Normal",
            "detector-ui面向检测结果消费：样件/批次列表、单件多面图像与缺陷框叠加、"
            "缺陷列表回放、复核状态展示。"
            "数据以静态JSON文件直读为主，不提前架设重型后端。"
            "展示字段与FR-D4对齐：样件ID、批次ID、检测面、缺陷类别、位置框、置信度、"
            "检测时间、模型版本、复核状态。",
        ),
        ("Heading 2", "9.2 前后端数据交互与接口契约"),
        (
            "Normal",
            "交互原则：文件直读优先；复核回写确有需要时再开轻量API。"
            "只读数据路径建议：GET /data/detect/jobs-index.json（列表索引）；"
            "GET /data/detect/jobs/{piece_id}.json（单件明细，含faces[]与defects[]）。"
            "可选复核写入：POST /api/review（口令校验；action含confirm/reject/relabel；草稿态，是否本期实现按复核工作量关单）。"
            "统一错误码遵循前后端契约（400/401/404/422/502等）。"
            "算法落盘格式与前端jobs格式之间，sync脚本为唯一转换入口。",
        ),
        ("Heading 2", "9.3 统计分析"),
        (
            "Normal",
            "在结构化结果具备piece_id、batch_id、by_class、timestamp、model_version的前提下，"
            "支持按批次聚合缺陷数量与类别分布、按模型版本对比、按复核状态统计待审/已确认/驳回。"
            "效率对照试验应采用同批样件与同一缺陷标准，将搬运、观察、记录、复核与报告计入全流程时间，"
            "结果可服务第10章效率对比表。统计口径不得把异常筛查与YOLO主检召回串行相乘后当作P0指标。",
        ),
        ("Heading 2", "9.4 管理BI"),
        (
            "Normal",
            "管理BI与detector-ui同属展示层、不同用途：BI承载课题管理信息，包括里程碑/五节点进度、"
            "风险台账、M币预算与流水、团队工作流、知识库入口与意见箱。"
            "部署形态建议静态站点（如Cloudflare Pages）；公开管理信息与企业检测数据分离，"
            "企业检测明细默认不进入公网。"
            "BI与仓库结构化数据（如.project-spec）通过sync保持单事实源，避免双份维护。",
        ),
        ("Heading 2", "9.5 联调与第三方试操检查点"),
        (
            "Normal",
            "联调最小闭环：离线模块对样例图产出result.json与overlay→执行同步脚本生成jobs数据→"
            "detector-ui可打开列表与叠加框→（可选）完成一条复核回写。"
            "第三方试操检查点与使用说明对齐：无产线联机环境下完成放图、跑检测、查看结果；"
            "CPU路径可运行；字段与目录符合第8章约定。",
        ),
    ]
    # insert before h10 in reverse
    for style, text in reversed(blocks):
        insert_paragraph_before(h10, text, style)

    # --- TOC entries for new ch9 (after toc line of ch8) ---
    toc8 = None
    for p in doc.paragraphs:
        style = p.style.name if p.style is not None else ""
        t = para_text(p)
        if "toc" in style.lower() and t.startswith("8 离线检测模块与软件接口"):
            toc8 = p
            break
    if toc8 is None:
        for p in doc.paragraphs:
            style = p.style.name if p.style is not None else ""
            t = para_text(p)
            if "toc" in style.lower() and "8 离线检测模块" in t:
                toc8 = p
                break

    if toc8 is not None:
        toc_blocks = [
            ("toc 1", "9 检测软件与BI详细设计"),
            ("toc 2", "9.1 检测结果展示（detector-ui）"),
            ("toc 2", "9.2 前后端数据交互与接口契约"),
            ("toc 2", "9.3 统计分析"),
            ("toc 2", "9.4 管理BI"),
            ("toc 2", "9.5 联调与第三方试操检查点"),
        ]
        anchor = toc8
        for style, text in toc_blocks:
            anchor = insert_paragraph_after(anchor, text, style)

    doc.save(str(REVISED_CLEAN))
    try:
        seed.unlink(missing_ok=True)
    except OSError:
        pass
    print("built clean revised:", REVISED_CLEAN)


def compare_with_word() -> None:
    import pythoncom
    import win32com.client

    pythoncom.CoInitialize()
    word = None
    d1 = d2 = dcomp = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        word.UserName = "算法软件组"
        word.UserInitials = "AS"

        # CompareDocuments returns a new document with revisions
        d1 = word.Documents.Open(str(ORIG), ReadOnly=True)
        d2 = word.Documents.Open(str(REVISED_CLEAN), ReadOnly=True)
        # Destination:=wdCompareDestinationNew (2)
        dcomp = word.CompareDocuments(
            OriginalDocument=d1,
            RevisedDocument=d2,
            Destination=2,
            CompareFormatting=False,
            CompareCaseChanges=True,
            CompareWhitespace=False,
            CompareTables=True,
            CompareHeaders=True,
            CompareFootnotes=True,
            CompareTextboxes=True,
            CompareFields=True,
            CompareComments=True,
            CompareMoves=True,
            RevisedAuthor="算法软件组",
            IgnoreAllComparisonWarnings=True,
        )
        tmp_out = TMP_DIR / "detail_plan_tracked.docx"
        if tmp_out.exists():
            try:
                tmp_out.unlink()
            except OSError:
                tmp_out = TMP_DIR / f"detail_plan_tracked_{int(time.time())}.docx"
        dcomp.SaveAs2(str(tmp_out))
        # copy into repo design folder
        if OUT_TRACKED.exists():
            try:
                OUT_TRACKED.unlink()
            except OSError:
                alt = OUT_DIR / f"03组_详细方案设计_算法软件修订_{int(time.time())}.docx"
                shutil.copy2(tmp_out, alt)
                print("tracked revisions saved (alt):", alt)
                return
        shutil.copy2(tmp_out, OUT_TRACKED)
        print("tracked revisions saved:", OUT_TRACKED)

        # Ensure revisions visible
        try:
            dcomp.ShowRevisions = True
            dcomp.TrackRevisions = False  # comparison doc already has revision marks
            dcomp.Save()
        except Exception as e:
            print("warn show revisions:", e)
    finally:
        for d in (dcomp, d2, d1):
            if d is not None:
                try:
                    d.Close(False)
                except Exception:
                    pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass


def main() -> int:
    if not ORIG.exists():
        print("missing original", ORIG)
        return 1
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    TMP_DIR.mkdir(parents=True, exist_ok=True)
    build_revised_clean()
    time.sleep(1)
    compare_with_word()
    target = OUT_TRACKED if OUT_TRACKED.exists() else None
    if target is None:
        candidates = sorted(OUT_DIR.glob("03组_详细方案设计_算法软件修订_*.docx"))
        target = candidates[-1] if candidates else None
    if target is None:
        print("ERR: tracked output missing")
        return 2
    doc = Document(str(target))
    texts = [para_text(p) for p in doc.paragraphs if para_text(p)]
    checks = [
        "3.1 总体架构原则",
        "8 离线检测模块与软件接口",
        "8.5 与前端的数据交接",
        "9 检测软件与BI详细设计",
        "10 测试情况与效率对比",
        "7.3 近线精修与工具边界",
    ]
    for c in checks:
        ok = any(c in t for t in texts)
        print(("OK" if ok else "MISSING"), c)
    print("FINAL", target)
    return 0


if __name__ == "__main__":
    sys.exit(main())
