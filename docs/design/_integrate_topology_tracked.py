# -*- coding: utf-8 -*-
"""Integrate software topology into design docx with Word Track Changes."""
from __future__ import annotations

import copy
import re
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from lxml import etree
import matplotlib.pyplot as plt
from matplotlib import font_manager
from matplotlib.patches import FancyBboxPatch

# Prefer Chinese-capable Windows fonts for topology labels
for _fname in ("Microsoft YaHei", "SimHei", "SimSun", "Arial Unicode MS"):
    try:
        font_manager.findfont(_fname, fallback_to_default=False)
        plt.rcParams["font.sans-serif"] = [_fname, "DejaVu Sans"]
        plt.rcParams["axes.unicode_minus"] = False
        break
    except Exception:
        continue

DOCX = Path(
    r"a:\Projects\MEM-AIVisdefect\docs\design"
    r"\03组_玻纤拉挤电芯压条_AI视觉检测_详细方案设计_编号目录整理版_算法软件修订.docx"
)
FIG_PATH = Path(r"a:\Projects\MEM-AIVisdefect\docs\design\_figs\software-topology.png")
AUTHOR = "算法软件组"
DATE = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def enable_track_revisions(doc: Document) -> None:
    settings = doc.settings.element
    if settings.find(qn("w:trackRevisions")) is None:
        settings.insert(0, OxmlElement("w:trackRevisions"))


def next_revision_id(doc: Document) -> int:
    body = doc.element.body
    ids = []
    for tag in ("w:ins", "w:del"):
        for el in body.findall(".//" + qn(tag)):
            rid = el.get(qn("w:id"))
            if rid is not None and str(rid).isdigit():
                ids.append(int(rid))
    return (max(ids) + 1) if ids else 1


class RevFactory:
    def __init__(self, start_id: int):
        self._id = start_id

    def _bump(self) -> str:
        i = self._id
        self._id += 1
        return str(i)

    def make_ins(self, text: str, *, rpr: etree._Element | None = None) -> etree._Element:
        ins = OxmlElement("w:ins")
        ins.set(qn("w:id"), self._bump())
        ins.set(qn("w:author"), AUTHOR)
        ins.set(qn("w:date"), DATE)
        run = OxmlElement("w:r")
        if rpr is not None:
            run.append(copy.deepcopy(rpr))
        t = OxmlElement("w:t")
        if text[:1].isspace() or text[-1:].isspace():
            t.set(qn("xml:space"), "preserve")
        t.text = text
        run.append(t)
        ins.append(run)
        return ins

    def make_del(self, text: str, *, rpr: etree._Element | None = None) -> etree._Element:
        dele = OxmlElement("w:del")
        dele.set(qn("w:id"), self._bump())
        dele.set(qn("w:author"), AUTHOR)
        dele.set(qn("w:date"), DATE)
        run = OxmlElement("w:r")
        if rpr is not None:
            run.append(copy.deepcopy(rpr))
        dt = OxmlElement("w:delText")
        if text[:1].isspace() or text[-1:].isspace():
            dt.set(qn("xml:space"), "preserve")
        dt.text = text
        run.append(dt)
        dele.append(run)
        return dele


def paragraph_plain_text(p_el: etree._Element) -> str:
    """Visible + inserted text (exclude delText)."""
    parts = []
    for node in p_el.iter():
        if node.tag == qn("w:delText"):
            continue
        if node.tag == qn("w:t") and node.text:
            # skip if ancestor is w:del
            parent = node.getparent()
            skip = False
            while parent is not None:
                if parent.tag == qn("w:del"):
                    skip = True
                    break
                parent = parent.getparent()
            if not skip:
                parts.append(node.text)
    return "".join(parts)


def clear_runs_keep_pPr(p_el: etree._Element) -> etree._Element | None:
    sample_rpr = None
    for child in list(p_el):
        tag = child.tag.split("}")[-1]
        if tag == "pPr":
            continue
        if sample_rpr is None:
            rpr = child.find(".//" + qn("w:rPr"))
            if rpr is not None:
                sample_rpr = rpr
        p_el.remove(child)
    return sample_rpr


def set_tracked_paragraph_text(
    p_el: etree._Element, rev: RevFactory, text: str, *, mark_old_as_del: bool = True
) -> None:
    old_text = paragraph_plain_text(p_el).strip()
    # also capture prior delText-only? no — only replace current non-deleted content
    rpr = clear_runs_keep_pPr(p_el)
    if mark_old_as_del and old_text and old_text != text:
        p_el.append(rev.make_del(old_text, rpr=rpr))
    if text:
        p_el.append(rev.make_ins(text, rpr=rpr))


def add_tracked_para_after(
    doc: Document,
    ref_element: etree._Element,
    rev: RevFactory,
    text: str,
    style_name: str,
) -> etree._Element:
    p = doc.add_paragraph(style=style_name)
    p_el = p._element
    clear_runs_keep_pPr(p_el)
    if text:
        p_el.append(rev.make_ins(text))
    ref_element.addnext(p_el)
    return p_el


def render_topology_figure(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fig, ax = plt.subplots(figsize=(11.5, 7.2), dpi=160)
    ax.set_xlim(0, 115)
    ax.set_ylim(0, 72)
    ax.axis("off")
    fig.patch.set_facecolor("white")

    def box(x, y, w, h, title, lines, edge="#333333", face="#F7F7F7", lw=1.2, ls="-"):
        patch = FancyBboxPatch(
            (x, y),
            w,
            h,
            boxstyle="round,pad=0.25,rounding_size=0.6",
            linewidth=lw,
            edgecolor=edge,
            facecolor=face,
            linestyle=ls,
        )
        ax.add_patch(patch)
        ax.text(
            x + w / 2,
            y + h - 0.9,
            title,
            ha="center",
            va="top",
            fontsize=9,
            fontweight="bold",
            color="#222",
        )
        for i, line in enumerate(lines):
            ax.text(
                x + 0.5,
                y + h - 2.1 - i * 0.95,
                line,
                ha="left",
                va="top",
                fontsize=7.5,
                color="#333",
            )

    box(
        2,
        58,
        28,
        12,
        "现场层（演进）",
        [
            "五段线：上线→姿态→环拍→分拣→分箱",
            "PLC / 传感器 / 触发（待需求表）",
            "气动推杆等执行器（OD-4.1 open）",
        ],
        edge="#888",
        face="#FAFAFA",
        ls="--",
    )
    box(
        34,
        58,
        28,
        12,
        "采集层（过程基线）",
        [
            "600万卷帘面阵 + 16mm 定焦",
            "光源 DOE 后冻结（OD-COM-05）",
            "停稳→拍摄→传图；OpenCV 标定",
        ],
        edge="#555",
        face="#F3F6FA",
    )
    box(
        66,
        58,
        46,
        12,
        "数据层（文件事实源）",
        [
            "raw-data 版本化原图 · 无 DB / MQ / MES",
            "标注 7:2:1 · MLflow 本地 mlruns",
            "result.json + overlay.jpg · .project-spec",
        ],
        edge="#2F6F9F",
        face="#EEF5FA",
    )
    box(
        2,
        34,
        36,
        20,
        "算法层 · 拆分部署（ADR-001）",
        [
            "P0 硬交付：离线 YOLO（yolo11s-seg）",
            "可选增强：EfficientAD（不进验收分母相乘）",
            "演进·在线 L1：PatchCore→可选蒸馏 YOLO",
            "演进·近线：Grounding DINO + SAM2（工具层）",
            "演进·飞轮 L4：VLM 预标注→人工确认 10%",
        ],
        edge="#1B7A4A",
        face="#EEF8F1",
        lw=1.8,
    )
    box(
        42,
        34,
        34,
        20,
        "应用层",
        [
            "detector-ui：结果框/列表/回放/复核",
            "BI：节点/风险/M币/知识库（CF Pages）",
            "薄后端：Pages Functions 意见箱写回",
            "契约：静态 JSON + 可选 review API",
        ],
        edge="#B07A00",
        face="#FFF8E8",
    )
    box(
        80,
        34,
        32,
        20,
        "协作侧车（非 AOI）",
        [
            "feishu-cursor-bridge",
            "飞书 <-> Cursor Agent",
            "不进入产线节拍闭环",
        ],
        edge="#888",
        face="#F5F5F5",
        ls="--",
    )
    box(
        2,
        8,
        110,
        20,
        "P0 验收主路径（最短闭环 · 不经 PLC / 在线 L1 / 近线工具）",
        [
            "放图 / 采集入库 → offline-detect CLI（conda huaduo · CPU 可跑 NFR-1）",
            "→ yolo11s-seg 必选推理（+ EfficientAD 可选）→ result.json + overlay.jpg",
            "→ npm run sync-data → detector-ui 复核 → 确认样本回流训练集 → 权重版本升级",
            "部署：工控机/工作站离线交付包；BI 另部 Cloudflare Pages（企业检测数据默认不公开）",
        ],
        edge="#1B7A4A",
        face="#F7FBF8",
        lw=1.6,
    )
    ax.annotate(
        "",
        xy=(34, 64),
        xytext=(30, 64),
        arrowprops=dict(arrowstyle="->", color="#666", lw=1.2, ls="--"),
    )
    ax.annotate(
        "",
        xy=(66, 64),
        xytext=(62, 64),
        arrowprops=dict(arrowstyle="->", color="#666", lw=1.2),
    )
    ax.annotate(
        "",
        xy=(20, 54),
        xytext=(50, 58),
        arrowprops=dict(arrowstyle="->", color="#444", lw=1.1),
    )
    ax.annotate(
        "",
        xy=(55, 54),
        xytext=(80, 58),
        arrowprops=dict(arrowstyle="->", color="#2F6F9F", lw=1.1),
    )
    ax.annotate(
        "",
        xy=(59, 44),
        xytext=(38, 44),
        arrowprops=dict(arrowstyle="->", color="#1B7A4A", lw=1.2),
    )
    ax.text(
        57.5,
        70.5,
        "MEM-AIVisdefect 全项目软件拓扑（现场 / 采集 / 算法 / 数据 / 应用）",
        ha="center",
        fontsize=11,
        fontweight="bold",
    )
    ax.text(
        57.5,
        1.5,
        "实线=本期交付或基线　虚线=演进/未关单　对齐 ADR-001 / 软件拓扑图.md",
        ha="center",
        fontsize=7.5,
        color="#666",
    )
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def renumber_old_captions(doc: Document, rev: RevFactory) -> list[str]:
    """Shift existing 图1–3 → 2–4 and 表6–22 → 7–23 before inserting new ones."""
    log = []
    fig_map = {
        "图 1  缺陷正例、良品与边界样本对照图": "图 2  缺陷正例、良品与边界样本对照图",
        "图1  缺陷正例、良品与边界样本对照图": "图 2  缺陷正例、良品与边界样本对照图",
        "图 2  成像工位总装图与检测面布置图": "图 3  成像工位总装图与检测面布置图",
        "图 3  离线检测模块界面原型": "图 4  离线检测模块界面原型",
    }
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        full = paragraph_plain_text(p._element).strip()
        if style == "Figure Caption CN":
            for old, new in fig_map.items():
                if full == old or old in full:
                    set_tracked_paragraph_text(p._element, rev, new, mark_old_as_del=True)
                    log.append(f"figure: {old} -> {new}")
                    break
        if style == "Table Caption CN":
            m = re.match(r"表\s*(\d+)\s+(.*)$", full)
            if m:
                n = int(m.group(1))
                title = m.group(2).strip()
                if 6 <= n <= 22:
                    new = f"表 {n + 1}  {title}"
                    set_tracked_paragraph_text(p._element, rev, new, mark_old_as_del=True)
                    log.append(f"table: 表 {n} -> 表 {n + 1}")
    return log


def find_heading8(doc: Document):
    for p in doc.paragraphs:
        if not p.style or p.style.name != "Heading 1":
            continue
        t = paragraph_plain_text(p._element).strip()
        if t in ("8", "8 ") or re.fullmatch(r"8\s*", t):
            return p
        if t.startswith("8 ") and "软件" not in t and len(t) < 8:
            return p
    return None


def fix_heading_8(doc: Document, rev: RevFactory) -> str:
    p = find_heading8(doc)
    if p is None:
        return "heading 8 not found"
    visible = paragraph_plain_text(p._element).strip()
    if "软件系统设计" in visible:
        return "heading 8 already named"
    if re.fullmatch(r"8\s*", visible):
        # keep existing number run if any; safest: del+ins whole title
        set_tracked_paragraph_text(p._element, rev, "8 软件系统设计", mark_old_as_del=True)
        return "heading 8 -> 8 软件系统设计 (tracked)"
    set_tracked_paragraph_text(p._element, rev, "8 软件系统设计", mark_old_as_del=True)
    return f"heading 8 replaced from {visible!r}"


def main() -> None:
    print("Rendering topology figure...")
    render_topology_figure(FIG_PATH)
    print("Figure:", FIG_PATH, "exists=", FIG_PATH.exists())

    print("Opening docx...")
    # backup first
    bak = DOCX.with_suffix(".docx.bak-before-topology")
    if not bak.exists():
        bak.write_bytes(DOCX.read_bytes())
        print("Backup:", bak)

    doc = Document(str(DOCX))
    enable_track_revisions(doc)
    rev = RevFactory(next_revision_id(doc))

    print("Renumbering old captions first...")
    for line in renumber_old_captions(doc, rev):
        print(" ", line)

    # Locate §3 placeholders by position after "3 总体方案设计"
    idx3 = None
    for i, p in enumerate(doc.paragraphs):
        if p.style and p.style.name == "Heading 1":
            if paragraph_plain_text(p._element).strip().endswith("总体方案设计"):
                idx3 = i
                break
    if idx3 is None:
        raise RuntimeError("section 3 not found")

    p31 = doc.paragraphs[idx3 + 1]
    p31_body = doc.paragraphs[idx3 + 2]
    p32 = doc.paragraphs[idx3 + 3]
    p32_body = doc.paragraphs[idx3 + 4]
    p_extra = doc.paragraphs[idx3 + 5]

    print("Filling §3.1 / §3.2 with tracked revisions...")
    set_tracked_paragraph_text(p31._element, rev, "3.1 总体架构与软件拓扑", mark_old_as_del=True)
    set_tracked_paragraph_text(
        p31_body._element,
        rev,
        (
            "总体架构遵循三条原则：（1）验收主线最短化——任务书 P0（≥3 类、召回≥80%、准确率≥85%、50 件）"
            "由离线检测模块一线承载；（2）部署形态按 ADR-001 修正——在线轻量、近线精修、离线飞轮拆分，"
            "禁止把四层能力串成产线每帧唯一判决链；（3）受控成像先于算法——光源经 DOE 冻结后再固化指标。"
            "全项目软件拓扑覆盖现场机械/PLC、采集、算法、数据、前后端应用与协作侧车；实线表示本期交付或过程基线，"
            "虚线表示演进方向或未关单事项。事实源为文件与 Git 快照，不引入工业数据库或消息队列。"
        ),
        mark_old_as_del=True,
    )

    fig_cap_el = add_tracked_para_after(
        doc, p31_body._element, rev, "图 1  全项目软件拓扑图", "Figure Caption CN"
    )
    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.add_run().add_picture(str(FIG_PATH), width=Cm(15.8))
    fig_cap_el.addnext(pic_p._element)

    note_el = add_tracked_para_after(
        doc,
        pic_p._element,
        rev,
        (
            "读图约定：现场 PLC/分拣、在线 L1、近线 DINO/SAM2、VLM 飞轮为演进能力；"
            "P0 验收仅经「放图→离线 YOLO→结构化结果/叠加图→detector-ui 复核→样本回流」。"
            "详细说明见仓库 docs/design/软件拓扑图.md。"
        ),
        "Normal",
    )

    set_tracked_paragraph_text(
        p32._element, rev, "3.2 分层边界、部署形态与验收主路径", mark_old_as_del=True
    )
    set_tracked_paragraph_text(
        p32_body._element,
        rev,
        (
            "分层边界：（1）现场层——五段线与 PLC/执行器，本期非硬交付，深联调与安全认证不在范围；"
            "（2）采集层——600 万卷帘面阵、16mm 定焦、走-停-走分段拍摄与 OpenCV 标定，分辨率/光源等见未决项；"
            "（3）算法层——离线 yolo11s-seg 为 P0 主检，EfficientAD 可选增强；在线 PatchCore、近线 DINO/SAM2、"
            "VLM 飞轮为演进工具，不进离线交付包硬依赖；（4）数据层——raw-data、标注集、result.json、"
            ".project-spec 经 npm run sync-data 同步；（5）应用层——detector-ui 面向检测复核，BI 面向项目管理，"
            "写接口仅意见箱等薄 Pages Functions；（6）协作侧车——飞书-Cursor 桥，不进入 AOI 运行时。"
            "本期部署：单台工控机/工作站 conda 环境 huaduo 承载离线模块；BI 部署于 Cloudflare Pages。"
            "P0 最短闭环不经 PLC、不经在线 L1、不经近线工具层。"
        ),
        mark_old_as_del=True,
    )

    # Ensure §3.2 follows note (it should already be after p31_body originally;
    # after inserts, order is p31, body, fig, pic, note, p32... — correct if p32 was after p31_body)
    # Re-anchor: move p32 block after note if needed
    if note_el.getnext() is not p32._element:
        # detach and place after note
        for el in (p_extra._element, p32_body._element, p32._element):
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)
        note_el.addnext(p32._element)
        p32._element.addnext(p32_body._element)
        p32_body._element.addnext(p_extra._element)

    tbl_cap_el = add_tracked_para_after(
        doc, p32_body._element, rev, "表 6  软件拓扑分层组件清单", "Table Caption CN"
    )

    table = doc.add_table(rows=7, cols=4)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    headers = ["层", "核心组件", "技术/接口要点", "本期状态"]
    rows = [
        ["现场", "五段线 / PLC / 推杆", "触发与分拣；流水线需求表", "演进（虚线）"],
        ["采集", "相机 / 光源 / 标定", "6MP 卷帘 + DOE + OpenCV", "过程基线"],
        ["算法", "离线 YOLO（主）", "yolo11s-seg CLI；CPU 可跑", "P0 硬交付"],
        ["数据", "文件 + Git", "raw / JSON / sync-data；无 DB·MQ", "已定"],
        ["应用", "detector-ui + BI", "React/Vite；CF Pages Functions", "BI 已落地"],
        ["协作", "feishu-cursor-bridge", "飞书 <-> Cursor Agent", "侧车"],
    ]
    for j, h in enumerate(headers):
        cell_p = table.rows[0].cells[j].paragraphs[0]
        clear_runs_keep_pPr(cell_p._element)
        cell_p._element.append(rev.make_ins(h))
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell_p = table.rows[i + 1].cells[j].paragraphs[0]
            clear_runs_keep_pPr(cell_p._element)
            cell_p._element.append(rev.make_ins(val))
    tbl_cap_el.addnext(table._element)

    set_tracked_paragraph_text(
        p_extra._element,
        rev,
        (
            "接口边摘要：离线模块 I/O 见《offline-module-interface》；"
            "detector-ui / BI 契约见《22-前后端API契约》；"
            "现场触发与连续/启停输送耦合见 OD-COM-03/04，关单前不承诺产线联机。"
        ),
        mark_old_as_del=True,
    )
    table._element.addnext(p_extra._element)

    print(fix_heading_8(doc, rev))

    doc.save(str(DOCX))
    print("Saved", DOCX)

    # Verify
    doc2 = Document(str(DOCX))
    tr = doc2.settings.element.find(qn("w:trackRevisions"))
    print("trackRevisions enabled:", tr is not None)
    body = doc2.element.body
    print("w:ins count:", len(body.findall(".//" + qn("w:ins"))))
    print("w:del count:", len(body.findall(".//" + qn("w:del"))))
    # author/date on recent inserts
    recent = [
        el
        for el in body.findall(".//" + qn("w:ins"))
        if el.get(qn("w:author")) == AUTHOR and el.get(qn("w:date")) == DATE
    ]
    print("new inserts this run:", len(recent))
    print("new deletes this run:", len([
        el for el in body.findall(".//" + qn("w:del"))
        if el.get(qn("w:author")) == AUTHOR and el.get(qn("w:date")) == DATE
    ]))

    print("--- §3 region preview ---")
    in_section = False
    for p in doc2.paragraphs:
        style = p.style.name if p.style else ""
        plain = paragraph_plain_text(p._element).strip()
        if style == "Heading 1" and "总体方案设计" in plain:
            in_section = True
        if in_section and style == "Heading 1" and "总体方案设计" not in plain:
            break
        if not in_section:
            continue
        has_blip = "a:blip" in etree.tostring(p._element, encoding="unicode")
        dels = [
            (t.text or "")
            for dele in p._element.findall(qn("w:del"))
            for t in dele.findall(".//" + qn("w:delText"))
        ]
        del_preview = "".join(dels)[:40]
        print(f"[{style}] plain={plain[:70]!r} del={del_preview!r} pic={has_blip}")


if __name__ == "__main__":
    main()
